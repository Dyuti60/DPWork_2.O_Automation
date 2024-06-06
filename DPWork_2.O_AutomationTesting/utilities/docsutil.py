import docx
import os
import sys
from docx.shared import Inches
from docx.shared import RGBColor
from exception import CustomException

def createDocxFile():
    try:
        document = docx.Document()
        return document
    except Exception as e:
        raise CustomException(e,sys)

def addMainHeading(document, documentTitle):
    try:
        document.add_heading(str(documentTitle),0)
    except Exception as e:
        raise CustomException(e,sys)

def addMediumHeading(document, documentTitle):
    try:
        document.add_heading(str(documentTitle),2)
    except Exception as e:
        raise CustomException(e,sys)

def addSmallHeading(document, documentTitle):
    try:
        document.add_heading(str(documentTitle),3)
    except Exception as e:
        raise CustomException(e,sys)

def appendContent(document, content):
    try:
        document.add_paragraph(str(content))
    except Exception as e:
        raise CustomException(e,sys)

def appendContentWithBlueColor(document, content):
    try:
        paragraph=document.add_paragraph().add_run(content)
        paragraph.font.color.rgb = RGBColor.from_string('0000FF')
    except Exception as e:
        raise CustomException(e,sys)

def appendContentWithFailColor(document, mainContent):
    try:
        paragraph = document.add_paragraph(str(mainContent)+" ")
        run = paragraph.add_run('Fail')
        run.font.color.rgb = RGBColor(255, 0, 0)
    except Exception as e:
        raise CustomException(e,sys)
    

def appendContentWithPassColor(document, mainContent):
    try:
        paragraph = document.add_paragraph(str(mainContent)+" ")
        run = paragraph.add_run('Pass')
        run.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
    except Exception as e:
        raise CustomException(e,sys)

def insertImageInDocx(document, image):
    try:
        p=document.add_paragraph()
        r=p.add_run()
        r.add_picture(image, width=Inches(6.5), height=Inches(3.5))
    except Exception as e:
        raise CustomException(e,sys)

def saveDocument(document, saveFilename):
    try:
       desired_location=os.getcwd()+"\\documentation\\"+saveFilename
       document.save(desired_location)
       ''' os.chdir("../")
        desiredLocation=os.getcwd()+'\\documentation\\dpWorkLoginPageDocumentation\\{}'.format(str(saveFilename))
        document.save("{}.docx".format(str(desiredLocation)))
        '''
    except Exception as e:
        raise CustomException(e,sys)

def createDedicatedSSFolder(foldername):
    try:
        screenshot_folderpath=os.getcwd()+"//screenshots//"+foldername
        os.makedirs(screenshot_folderpath,exist_ok=True)
        location=screenshot_folderpath+'//'
        return location
    except Exception as e:
        raise CustomException(e,sys)

def takeAndSaveScreenshotUnique(location,username,serialNumber,driver):
    try:
        final_location=location+'\\SS-{} - {}.png'.format(serialNumber,username)
        driver.save_screenshot(final_location)
        return final_location
    except Exception as e:
        raise CustomException(e,sys)

def takeAndSaveScreenshotCommon(location,username,serialNumber,driver):
    try:
        final_location=location+'\\SS-{}-LoginTest_DpWorker_UserName - {}.png'.format(serialNumber,username)
        driver.save_screenshot(final_location)
    except Exception as e:
        raise CustomException(e,sys)

def collectScreenshot(screenshotFolder,username):
    try:
        location=os.getcwd()+"//{}".format(screenshotFolder)
        desired_screenshots=[]
        for ss in os.listdir(location):
            if ss.__contains__(username):
                desired_screenshots.append(location+"//{}".format(ss))
        return desired_screenshots
    except Exception as e:
        raise CustomException(e, sys)

def deleteScreenshots(screenshotFolder,username):
    try:
        location=os.getcwd()+"//{}".format(screenshotFolder)
        desired_screenshots=[]
        for ss in os.listdir(location):
            if ss.__contains__(username):
                os.remove(location+"//{}".format(ss))
    except Exception as e:
        raise CustomException(e,sys)